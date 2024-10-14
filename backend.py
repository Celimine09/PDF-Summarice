import os
import PyPDF2
from transformers import T5ForConditionalGeneration, T5Tokenizer
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

model_path = r'results_t5bbc_billsum_resume\\results_t5bbc_billsum_resume\\checkpoint-130'
model = T5ForConditionalGeneration.from_pretrained(model_path)
tokenizer = T5Tokenizer.from_pretrained('results_t5bbc_billsum_resume\\results_t5bbc_billsum_resume')

def summarize_text_bart(text):
    inputs = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=1024, truncation=True)
    outputs = model.generate(
        inputs, 
        max_length=300, 
        min_length=100, 
        length_penalty=1.2, 
        num_beams=5, 
        early_stopping=True
    )
    summary = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return summary

def summarize_pdf_bart(file_path):
    try:
        with open(file_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()

            summary = summarize_text_bart(text)
            return summary, text
    except Exception as e:
        print(f"Error processing PDF file {file_path}: {e}")
        raise e

def calculate_similarities(texts):
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(texts)
    cosine_similarities = cosine_similarity(tfidf_matrix)
    return cosine_similarities * 100

def process_pdfs_in_folder_bart(folder_path, output_docx_path, progress_callback=None):
    summaries = []
    full_texts = []
    pdf_files = [f for f in os.listdir(folder_path) if f.endswith('.pdf')]
    total_files = len(pdf_files)

    for i, file_name in enumerate(pdf_files):
        file_path = os.path.join(folder_path, file_name)
        print(f"Processing {file_name}...")

        try:
            summary, full_text = summarize_pdf_bart(file_path)
            summaries.append((summary, file_name))
            full_texts.append(full_text)

            if progress_callback:
                progress_callback(int((i + 1) / total_files * 100))  # Update progress

        except Exception as e:
            print(f"Failed to process file {file_name}: {e}")
            continue  # Skip this file if it fails

    # Calculate similarities
    similarities = calculate_similarities(full_texts)
    average_similarities = np.mean(similarities, axis=1)

    # Combine summaries with similarities
    summaries_with_scores = [(summary, score, file_name) for (summary, file_name), score in zip(summaries, average_similarities)]

    # Sort summaries by similarity score
    summaries_with_scores.sort(key=lambda x: x[1], reverse=True)

    # Create DOC file
    doc = Document()
    doc.add_heading(os.path.basename(folder_path), 0)

    showgui = []
    for idx, (summary, score, file_name) in enumerate(summaries_with_scores):
        file_name_without_ext = os.path.splitext(file_name)[0]
        if idx < 5:
            showgui.append((file_name_without_ext, score))
        doc.add_heading(f'{file_name_without_ext} | Score : {score:.4f}', level=1)
        doc.add_paragraph(summary)

    doc.save(output_docx_path)
    print(f"Output saved to {output_docx_path}")
    return showgui

if __name__ == "__main__":
    folder_path = "CHEF"
    output_docx_path = "summary.docx"
    process_pdfs_in_folder_bart(folder_path, output_docx_path)